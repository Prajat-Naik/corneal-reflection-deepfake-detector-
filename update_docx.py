import os
import re
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_docx():
    doc = docx.Document()
    
    # Standard margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Read the polished markdown source
    with open('IEEE_Research_Paper.md', 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    
    # Custom heading helper
    def add_custom_heading(text, level, space_before=12, space_after=6):
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(space_before)
        heading.paragraph_format.space_after = Pt(space_after)
        heading.paragraph_format.keep_with_next = True
        
        run = heading.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(14)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)
            run.italic = True
        return heading

    # Document Header/Title Block
    doc.add_paragraph()  # Spacing
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Detection of deepfake images using corneal reflection")
    run_title.bold = True
    run_title.font.size = Pt(18)
    p_title.paragraph_format.space_after = Pt(12)

    # Author Affiliations
    p_auth = doc.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_auth1 = p_auth.add_run("1st Prajat Anand Naik\nUSN: 4MT24MC064\nDepartment of Master of Computer Applications\nMangalore Institute of Technology and Engineering\nMangaluru, India\n\n")
    r_auth1.font.size = Pt(10)
    
    r_guide = p_auth.add_run("2nd Shylesh B C (Project Guide)\nDepartment of Master of Computer Applications\nMangalore Institute of Technology and Engineering\nMangaluru, India")
    r_guide.font.size = Pt(10)
    p_auth.paragraph_format.space_after = Pt(18)

    # Iterative line-by-line parsing
    i = 0
    in_abstract = False
    
    while i < len(lines):
        line = lines[i].strip()
        i += 1
        
        if not line:
            continue
            
        # Ignore horizontal rules
        if line == '---':
            continue
            
        # Parse markdown headings
        if line.startswith('# '):
            continue  # Title already added manually
        elif line.startswith('## '):
            title_text = line[3:].strip()
            if title_text.upper() == 'ABSTRACT':
                add_custom_heading('ABSTRACT', 2, 18, 6)
                in_abstract = True
            elif title_text.upper() == 'KEYWORDS':
                add_custom_heading('KEYWORDS', 2, 12, 4)
                in_abstract = False
            else:
                add_custom_heading(title_text, 1, 18, 8)
        elif line.startswith('### '):
            title_text = line[4:].strip()
            add_custom_heading(title_text, 2, 12, 6)
        elif line.startswith('#### '):
            title_text = line[5:].strip()
            add_custom_heading(title_text, 3, 10, 4)
        elif line.startswith('* ') or line.startswith('- '):
            # Bullet list items
            li = doc.add_paragraph(style='List Bullet')
            li.paragraph_format.space_after = Pt(4)
            text = line[2:].strip()
            # Parse bold runs
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = li.add_run(part[2:-2])
                    run.bold = True
                else:
                    li.add_run(part)
        elif re.match(r'^\d+\.\s+', line):
            # Numbered list items
            match = re.match(r'^(\d+)\.\s+(.*)', line)
            num = match.group(1)
            text = match.group(2).strip()
            
            li = doc.add_paragraph()
            li.paragraph_format.left_indent = Inches(0.25)
            li.paragraph_format.space_after = Pt(4)
            
            run_num = li.add_run(f"{num}. ")
            run_num.bold = True
            
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = li.add_run(part[2:-2])
                    run.bold = True
                else:
                    li.add_run(part)
        elif line.startswith('>'):
            # Citations/Placeholders
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            clean_line = line.lstrip('>').strip()
            run = p.add_run(clean_line)
            run.italic = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = docx.shared.RGBColor(100, 100, 100)
        else:
            # Paragraph text
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if in_abstract:
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(line)
                run.italic = True
            else:
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    else:
                        p.add_run(part)
                        
    doc.save('4MT24MC064.docx')
    print("Word document built successfully!")

if __name__ == '__main__':
    build_docx()
